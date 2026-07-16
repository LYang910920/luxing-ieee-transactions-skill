"""Safe local text extraction.

The standard-library core reads text-like files. PDF and DOCX support is
optional and never falls back to OCR.
"""

from __future__ import annotations

from html.parser import HTMLParser
from pathlib import Path
from typing import Iterable

TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".tex", ".rst", ".bib", ".csv", ".yaml", ".yml"}
SUPPORTED_SUFFIXES = TEXT_SUFFIXES | {".html", ".htm", ".pdf", ".docx"}


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag.lower() in {"script", "style", "noscript"}:
            self._skip_depth += 1
        elif tag.lower() in {"p", "br", "div", "section", "article", "h1", "h2", "h3", "li"}:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() in {"script", "style", "noscript"} and self._skip_depth:
            self._skip_depth -= 1
        elif tag.lower() in {"p", "div", "section", "article", "li"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self._skip_depth:
            self.parts.append(data)

    def text(self) -> str:
        return "".join(self.parts)


def read_text(path: str | Path) -> str:
    """Read a supported local document and return extracted text.

    Raises a clear RuntimeError when an optional document dependency is not
    installed. OCR is intentionally not used.
    """

    p = Path(path)
    if not p.exists() or not p.is_file():
        raise FileNotFoundError(f"Document not found: {p}")

    suffix = p.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return p.read_text(encoding="utf-8", errors="replace")

    if suffix in {".html", ".htm"}:
        parser = _HTMLTextExtractor()
        parser.feed(p.read_text(encoding="utf-8", errors="replace"))
        return parser.text()

    if suffix == ".pdf":
        try:
            import fitz  # type: ignore
        except ImportError as exc:
            raise RuntimeError(
                "PDF extraction requires PyMuPDF. Install with: "
                "python -m pip install -e '.[documents]'"
            ) from exc
        doc = fitz.open(p)
        try:
            return "\n\n".join(page.get_text("text") for page in doc)
        finally:
            doc.close()

    if suffix == ".docx":
        try:
            from docx import Document  # type: ignore
        except ImportError as exc:
            raise RuntimeError(
                "DOCX extraction requires python-docx. Install with: "
                "python -m pip install -e '.[documents]'"
            ) from exc
        doc = Document(p)
        paragraphs = [paragraph.text for paragraph in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                paragraphs.append("\t".join(cell.text for cell in row.cells))
        return "\n\n".join(paragraphs)

    raise ValueError(f"Unsupported document type: {suffix or '<no suffix>'}")


def iter_documents(path: str | Path, recursive: bool = True) -> Iterable[Path]:
    """Yield supported files in deterministic order."""

    p = Path(path)
    if p.is_file():
        if p.suffix.lower() in SUPPORTED_SUFFIXES:
            yield p
        return
    if not p.exists():
        raise FileNotFoundError(f"Path not found: {p}")
    pattern = "**/*" if recursive else "*"
    for candidate in sorted(p.glob(pattern)):
        if candidate.is_file() and candidate.suffix.lower() in SUPPORTED_SUFFIXES:
            yield candidate
